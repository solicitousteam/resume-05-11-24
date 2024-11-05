# Fully Working
import os
import tempfile
import logging
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import StreamingResponse
import re
from io import BytesIO
from docx import Document
from docx.shared import Inches
import aspose.pdf as pdf

# Load SpaCy model

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

app = FastAPI()

# Regex patterns for sensitive information
EMAIL_PATTERN = re.compile(r'[a-zA-Z0-9._%+-]+@gmail\.com')
GITHUB_PATTERN = re.compile(r'https?://(www\.)?github\.com/[^\s]+', re.IGNORECASE)
LINKEDIN_PATTERN = re.compile(r'https?://(www\.)?linkedin\.com/[^\s]+', re.IGNORECASE)
MOBILE_PATTERN = re.compile(r'\b(\+?\d{1,3}[-.\s]?(\(?\d{1,4}?\)?[-.\s]?)?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9})\b')


def pdf_to_docx_aspose(pdf_path, docx_output_path):
    """Convert a PDF to DOCX using Aspose.PDF."""
    doc = pdf.Document(pdf_path)
    doc.save(docx_output_path, pdf.SaveFormat.DOC_X)
    logging.info(f"PDF has been converted to DOCX: {docx_output_path}")


def remove_sensitive_info(document):
    """Remove sensitive information like email, GitHub, LinkedIn, and mobile numbers."""
    for para in document.paragraphs:
        if EMAIL_PATTERN.search(para.text):
            logging.info(f"Removing email: {para.text}")
            para.text = EMAIL_PATTERN.sub("", para.text)
        if GITHUB_PATTERN.search(para.text):
            logging.info(f"Removing GitHub link: {para.text}")
            para.text = GITHUB_PATTERN.sub("", para.text).strip()
        if LINKEDIN_PATTERN.search(para.text):
            logging.info(f"Removing LinkedIn link: {para.text}")
            para.text = LINKEDIN_PATTERN.sub("", para.text).strip()
        if MOBILE_PATTERN.search(para.text):
            logging.info(f"Removing mobile number: {para.text}")
            para.text = MOBILE_PATTERN.sub("", para.text)


# def remove_names(document):
#     """Remove names from the document using SpaCy."""
#     for para in document.paragraphs:
#         doc = nlp(para.text)  # Process the paragraph with SpaCy
#         modified_text = " ".join([token.text for token in doc if token.ent_type_ != "PERSON"])
#         if modified_text != para.text:
#             logging.info(f"Removing name from paragraph: {para.text}")
#             para.text = modified_text


def insert_logo(document, logo_file):
    """Insert the logo image on the right side of the document header."""
    logo_added = False
    for section in document.sections:
        if not logo_added:
            header = section.header
            header_paragraph = header.paragraphs[0]
            run = header_paragraph.add_run()
            run.add_picture(logo_file, width=Inches(0.5), height=Inches(0.5))
            header_paragraph.alignment = 2  # Align right
            logo_added = True


def modify_docx(file_content: bytes, logo_content: bytes):
    """Process the DOCX file, removing sensitive information and adding a logo."""
    try:
        logging.info("Modifying DOCX file")
        document = Document(BytesIO(file_content))
        insert_logo(document, BytesIO(logo_content))
        remove_sensitive_info(document)
        # remove_names(document)

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output
    except Exception as e:
        logging.error(f"Error modifying DOCX: {e}")
        raise HTTPException(status_code=500, detail="Error processing DOCX file")


@app.post("/upload/")
async def upload_and_process_file(file: UploadFile = File(...), logo: UploadFile = File(...)):
    """Handle file upload, convert PDF if necessary, process, and return the modified DOCX."""
    logging.info(f"Received file: {file.filename} and logo: {logo.filename}")

    # Validate logo file types
    if not logo.filename.endswith(('.png', '.jpg', '.jpeg')):
        logging.warning(f"Unsupported logo format: {logo.filename}")
        raise HTTPException(status_code=400, detail="Unsupported logo format")

    # Read logo file content
    logo_content = await logo.read()

    # Read file content based on file type
    if file.filename.endswith('.pdf'):
        logging.info(f"Converting PDF to DOCX: {file.filename}")

        # Use tempfile for cross-platform compatibility
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_temp_file:
            pdf_temp_file.write(await file.read())
            pdf_temp_file.flush()  # Ensure all data is written
            pdf_temp_file_path = pdf_temp_file.name

        try:
            # Create a separate temporary DOCX file
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as docx_temp_file:
                docx_temp_file_path = docx_temp_file.name

            # Convert PDF to DOCX
            pdf_to_docx_aspose(pdf_temp_file_path, docx_temp_file_path)

            # Process the converted DOCX
            with open(docx_temp_file_path, "rb") as docx_file:
                docx_content = docx_file.read()
                modified_docx = modify_docx(docx_content, logo_content)

        finally:
            # Clean up temporary files
            os.remove(pdf_temp_file_path)
            os.remove(docx_temp_file_path)

    elif file.filename.endswith('.docx'):
        file_content = await file.read()
        modified_docx = modify_docx(file_content, logo_content)
    else:
        logging.warning(f"Unsupported file format: {file.filename}")
        raise HTTPException(status_code=400, detail="Unsupported file format")

    # Return the modified DOCX file
    return StreamingResponse(modified_docx,
                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={
                                 "Content-Disposition": f"attachment; filename=updated_{file.filename.replace('.pdf', '.docx')}"})
